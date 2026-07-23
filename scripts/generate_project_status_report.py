from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Informe_Estado_Web_John_Tours_2026.docx"
LOGO = ROOT / "frontend" / "public" / "john-tours-logo-cropped.png"

NAVY = "082447"
BLUE = "0F4C81"
GREEN = "0F7A4F"
GOLD = "D9A122"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF8F2"
PALE_GOLD = "FFF7DF"
GRAY = "52677B"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_fill(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    props = cell._tc.get_or_add_tcPr()
    width = props.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        props.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    props = table._tbl.tblPr
    table_width = props.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        props.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = props.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        props.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=NAVY, bold=False, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text: str, *, bold=False, color=NAVY, size=11, after=6, align=None, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if align is not None:
        paragraph.alignment = align
    set_run(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text: str, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run(paragraph.add_run(text), size=10.5, color=NAVY)
    return paragraph


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run(paragraph.add_run(text), size=10.5, color=NAVY)
    return paragraph


def add_callout(doc, title: str, text: str, fill=PALE_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.margin_top = 120
    cell.margin_bottom = 120
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    set_run(paragraph.add_run(title + "  "), size=10.5, color=accent, bold=True)
    set_run(paragraph.add_run(text), size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_status_table(doc):
    rows = [
        ("Experiencia pública", "Implementado", "Inicio, catálogo, detalle, confianza, publicidad y diseño adaptable."),
        ("Reserva y demostración", "Implementado", "Separación S/ 200, QR, código único, confirmación y modo sin cobro."),
        ("Contenido posterior al pago", "Implementado", "Guías PDF por destino, servicios extra y cita con mensaje formal."),
        ("Administración", "Implementado base", "Tours, reservas, pagos, configuración empresarial y políticas."),
        ("Pagos y correo reales", "Preparado", "Arquitectura lista; faltan credenciales, pruebas y configuración empresarial."),
        ("Dominio, hosting y correos", "Pendiente", "Se implementará cuando la empresa defina proveedor, dominio y cuentas."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 1700, 5360])
    for index, title in enumerate(("Área", "Estado", "Resumen")):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, NAVY)
        paragraph = cell.paragraphs[0]
        set_run(paragraph.add_run(title), size=10, color=WHITE, bold=True)
    for area, status, summary in rows:
        cells = table.add_row().cells
        values = (area, status, summary)
        status_color = GREEN if status == "Implementado" else GOLD if status in ("Preparado", "Implementado base") else GRAY
        for index, value in enumerate(values):
            if len(table.rows) % 2 == 1:
                set_cell_fill(cells[index], LIGHT_GRAY)
            paragraph = cells[index].paragraphs[0]
            set_run(paragraph.add_run(value), size=9.4, color=status_color if index == 1 else NAVY, bold=index == 1)
    set_table_geometry(table, [2300, 1700, 5360])


def add_heading(doc, text: str, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("JOHN TOURS PERÚ  |  INFORME DE ESTADO"), size=8, color=GRAY, bold=True)
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer_paragraph.add_run("Documento de trabajo - Julio 2026  |  Página "), size=8, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer_paragraph._p.append(field)


def build_report():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    if LOGO.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.add_run().add_picture(str(LOGO), width=Inches(2.65))
    add_paragraph(doc, "INFORME DE ESTADO Y HOJA DE RUTA", bold=True, color=GOLD, size=10, after=2)
    add_paragraph(doc, "Plataforma web John Tours Perú", bold=True, color=NAVY, size=28, after=5)
    add_paragraph(doc, "Funciones implementadas, preparación para producción y pendientes de dominio, hosting y correos corporativos", color=GRAY, size=13, after=18)

    metadata = doc.add_table(rows=2, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(metadata, [4680, 4680])
    entries = (("Fecha de actualización", "22 de julio de 2026"), ("Estado general", "Demostración funcional / preparación para producción"), ("Proyecto", "Agencia de Tours - John Tours Perú"), ("Próxima etapa", "Infraestructura y datos empresariales reales"))
    for index, (label, value) in enumerate(entries):
        cell = metadata.cell(index // 2, index % 2)
        set_cell_fill(cell, PALE_BLUE if index % 2 == 0 else PALE_GREEN)
        paragraph = cell.paragraphs[0]
        set_run(paragraph.add_run(label.upper() + "\n"), size=8, color=GRAY, bold=True)
        set_run(paragraph.add_run(value), size=10.2, color=NAVY, bold=True)
    doc.add_paragraph()

    add_callout(doc, "RESUMEN EJECUTIVO", "La plataforma ya permite promocionar tours, orientar al cliente, simular y registrar reservas, preparar pagos, desbloquear contenido posterior a la separación y coordinar citas. Para operar comercialmente faltan datos legales definitivos, pagos reales verificados, dominio, hosting, base de datos productiva y correos corporativos.")

    add_heading(doc, "1. Estado general", 1)
    add_status_table(doc)

    add_heading(doc, "2. Funciones implementadas", 1)
    add_heading(doc, "2.1 Experiencia comercial y publicidad", 2)
    for item in (
        "Inicio de alto impacto con identidad azul, verde y dorada, buscador y llamadas a la acción.",
        "Catálogo adaptable con filtros, precios, cupos, duración, moneda y destinos destacados.",
        "Secciones de confianza, acompañamiento, experiencias exclusivas, testimonios, preguntas frecuentes, historia y redes sociales.",
        "WhatsApp flotante con logo, disponibilidad del asesor, mensajes dinámicos y orientación según presupuesto, viajeros, fechas y destino.",
        "Diseño adaptable para celulares, tablets y computadoras con animaciones moderadas y navegación accesible.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.2 Reserva, separación y pago", 2)
    for item in (
        "Formulario de reserva con datos del pasajero, fecha del viaje y cantidad de personas.",
        "Separación referencial de S/ 200, QR informativo y código único por reserva.",
        "Envío del comprobante por WhatsApp para validación humana; no se confirma automáticamente sin revisión.",
        "Integración preparada para Culqi y Yape, con cálculo del monto en backend y protección de claves privadas.",
        "Modo demostración que permite presentar el recorrido completo sin realizar cobros.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.3 Beneficios después del pago", 2)
    for item in (
        "Guías PDF personalizadas con logo, imagen referencial y extras para Cusco, Orlando, Oxapampa, Ica y Egipto.",
        "Documento general para destinos nuevos que todavía no cuentan con una guía específica.",
        "Planificador de cita con fecha, hora, modalidad y tema de atención.",
        "Mensaje formal preparado con reserva, código de separación, monto y solicitud de adjuntar boleta o comprobante.",
        "En modo prueba el mensaje solo se visualiza o copia; nunca se envía ni abre WhatsApp automáticamente.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.4 Administración y backend", 2)
    for item in (
        "Panel para administrar tours, precios, monedas, cupos, imágenes, itinerarios e inclusiones.",
        "Consulta de reservas, pagos y configuración empresarial.",
        "Backend con Express, TypeScript, Prisma y MySQL; validación con esquemas y manejo centralizado de errores.",
        "Servicios separados para autenticación, reservas, pagos, correo, contacto y operaciones.",
        "Pruebas automáticas para políticas, testimonios, capacidad de salidas y adelantos.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "3. Recorrido actual del cliente", 1)
    for step in (
        "Descubre la marca y consulta opciones según su presupuesto.",
        "Filtra destinos y revisa el detalle completo del paquete.",
        "Registra sus datos y crea una reserva pendiente.",
        "Recibe un código único y las instrucciones para separar el cupo.",
        "Envía el comprobante para validación manual.",
        "Accede a la confirmación, los extras y la guía PDF del destino.",
        "Prepara una cita formal con un asesor para revisar itinerario, documentos y saldo.",
    ):
        add_number(doc, step)

    add_heading(doc, "4. Seguridad y límites actuales", 1)
    add_callout(doc, "REGLA DE PRODUCCIÓN", "La demostración no debe confundirse con un cobro real. Antes de aceptar dinero se deben configurar credenciales empresariales, políticas publicadas, webhook, auditoría y pruebas completas.", fill=PALE_GOLD, accent=GOLD)
    for item in (
        "Las claves privadas de Culqi solo deben existir en el backend y en variables seguras del hosting.",
        "No se guardarán claves de Yape, códigos SMS, contraseñas bancarias ni datos de tarjetas.",
        "El backend debe recalcular importes desde la reserva y registrar cada cambio de estado.",
        "Las fotografías y testimonios finales deben contar con autorización verificable.",
        "Los mensajes de WhatsApp se preparan, pero el cliente decide cuándo enviarlos.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "5. Implementación futura: dominio, hosting y correos", 1)
    add_paragraph(doc, "Esta fase está prevista, pero no debe activarse hasta que John Tours entregue los datos oficiales y seleccione proveedores. La web conservará el diseño actual y sustituirá progresivamente los datos demostrativos.", color=GRAY)

    add_heading(doc, "5.1 Dominio", 2)
    for item in (
        "Definir y registrar el dominio oficial de la empresa.",
        "Configurar DNS, HTTPS, redirección canónica y renovación automática.",
        "Conectar el dominio al frontend, API y plataforma de correo sin exponer claves.",
        "Actualizar enlaces, metadata, sitemap, redes sociales y documentos legales con la URL definitiva.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "5.2 Hosting e infraestructura", 2)
    for item in (
        "Publicar el frontend en Vercel u otra plataforma compatible con Next.js.",
        "Publicar la API en Railway, Render o un servidor administrado y conectar MySQL productivo.",
        "Configurar variables de entorno, copias de seguridad, monitoreo, logs y alertas.",
        "Separar ambientes de demostración, pruebas y producción.",
        "Activar CDN, compresión de imágenes, límites de tráfico y recuperación ante fallos.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "5.3 Correos corporativos", 2)
    add_paragraph(doc, "Cuentas sugeridas cuando exista el dominio oficial:", bold=True, color=NAVY, after=5)
    for item in (
        "reservas@dominio-oficial: confirmaciones, citas y documentación del pasajero.",
        "ventas@dominio-oficial: cotizaciones y seguimiento comercial.",
        "soporte@dominio-oficial: asistencia antes, durante y después del viaje.",
        "administracion@dominio-oficial: pagos, comprobantes y coordinación interna.",
        "no-reply@dominio-oficial: notificaciones automáticas sin recepción de respuestas.",
    ):
        add_bullet(doc, item)
    add_paragraph(doc, "La configuración deberá incluir SPF, DKIM y DMARC, SMTP seguro, plantillas con identidad visual y pruebas de entregabilidad. Estos correos se mostrarán en la web únicamente cuando las cuentas existan y hayan sido verificadas.", color=GRAY)

    doc.add_page_break()
    add_heading(doc, "6. Pendientes antes de publicar", 1)
    priorities = (
        ("Prioridad alta", "Razón social, RUC, dominio, políticas, datos de pago, precios, fechas, cupos e imágenes autorizadas."),
        ("Prioridad alta", "Credenciales reales de Culqi/Yape, webhook, validación de pagos y pruebas transaccionales."),
        ("Prioridad alta", "Base de datos productiva, copias de seguridad, cambio de credenciales demo y protección del panel administrativo."),
        ("Prioridad media", "Correos corporativos, SMTP, plantillas automáticas y seguimiento de entregabilidad."),
        ("Prioridad media", "SEO técnico, analítica, monitoreo de errores, rendimiento y accesibilidad."),
        ("Prioridad media", "Libro de Reclamaciones, términos, privacidad, cancelaciones y reembolsos publicados."),
        ("Mejora futura", "Calendario real de citas, recordatorios automáticos y registro del asesor asignado."),
        ("Mejora futura", "Carga de comprobantes dentro de la web y auditoría completa de aprobación o rechazo."),
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2100, 7260])
    for index, label in enumerate(("Nivel", "Implementación pendiente")):
        set_cell_fill(table.rows[0].cells[index], NAVY)
        set_run(table.rows[0].cells[index].paragraphs[0].add_run(label), size=10, color=WHITE, bold=True)
    for level, task in priorities:
        cells = table.add_row().cells
        fill = PALE_GOLD if level == "Prioridad alta" else PALE_BLUE if level == "Prioridad media" else PALE_GREEN
        set_cell_fill(cells[0], fill)
        set_run(cells[0].paragraphs[0].add_run(level), size=9.5, color=GOLD if level == "Prioridad alta" else BLUE if level == "Prioridad media" else GREEN, bold=True)
        set_run(cells[1].paragraphs[0].add_run(task), size=9.5, color=NAVY)
    set_table_geometry(table, [2100, 7260])

    add_heading(doc, "7. Información que debe entregar la empresa", 1)
    for item in (
        "Nombre de dominio elegido y responsable de su renovación.",
        "Razón social, RUC, dirección, teléfonos y representantes autorizados.",
        "Número, titular y QR empresarial de Yape; cuenta Culqi propia.",
        "Lista definitiva de correos corporativos y responsables de cada bandeja.",
        "Tours reales, precios, monedas, fechas, cupos, condiciones e inclusiones.",
        "Políticas aprobadas de privacidad, cambios, cancelación y reembolso.",
        "Fotografías, videos, testimonios y permisos de publicación.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "8. Criterios para salir a producción", 1)
    for item in (
        "Dominio y certificado HTTPS activos.",
        "Frontend, API y MySQL desplegados con copias de seguridad.",
        "Correos corporativos verificados y plantillas probadas.",
        "Pagos de prueba aprobados, webhook activo y conciliación verificada.",
        "Políticas legales publicadas y datos empresariales completos.",
        "Pruebas en móvil, escritorio, navegadores, accesibilidad y rendimiento.",
        "Monitoreo, analítica, alertas y procedimiento de soporte definidos.",
    ):
        add_bullet(doc, item)

    add_callout(doc, "RECOMENDACIÓN FINAL", "Mantener la web en modo demostración hasta completar las prioridades altas. El siguiente avance recomendado es definir dominio y datos empresariales; después se configuran hosting, base de datos, pagos reales y correos corporativos en un ambiente de pruebas antes de publicar.", fill=PALE_BLUE, accent=BLUE)

    doc.core_properties.title = "Informe de estado y hoja de ruta - John Tours Perú"
    doc.core_properties.subject = "Funciones implementadas y pendientes para producción"
    doc.core_properties.author = "John Tours Perú"
    doc.core_properties.keywords = "John Tours, web, dominio, hosting, correos corporativos, producción"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
